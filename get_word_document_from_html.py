from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt
from io import BytesIO

def generate_word_document(content):
    soup = BeautifulSoup(content, 'html.parser')

    doc = Document()

    def add_html_element_to_doc(element, doc):
        if not element.get_text(strip=True):
            return

        if element.name == 'h1':
            p = doc.add_heading(level=1).add_run(element.get_text())
            p.font.size = Pt(24)
            p.bold = True
        elif element.name == 'h2':
            p = doc.add_heading(level=2).add_run(element.get_text())
            p.font.size = Pt(20)
            p.bold = True
        elif element.name == 'h3':
            p = doc.add_heading(level=3).add_run(element.get_text())
            p.font.size = Pt(16)
            p.bold = True
        elif element.name == 'p':
            p = doc.add_paragraph(element.get_text())
        elif element.name == 'ul':
            for li in element.find_all('li'):
                p = doc.add_paragraph(li.get_text(), style='ListBullet')
        elif element.name == 'ol':
            for li in element.find_all('li'):
                p = doc.add_paragraph(li.get_text(), style='ListNumber')
        elif element.name == 'strong':
            run = doc.add_paragraph().add_run(element.get_text())
            run.bold = True
        elif element.name == 'em':
            run = doc.add_paragraph().add_run(element.get_text())
            run.italic = True
        elif element:
            p = doc.add_paragraph(element.get_text())

    for element in soup.contents:
        add_html_element_to_doc(element, doc)

    byte_stream = BytesIO()
    doc.save(byte_stream)

    byte_array = byte_stream.getvalue()
    byte_stream.close()

    return byte_array
